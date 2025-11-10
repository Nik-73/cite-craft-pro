#!/usr/bin/env python3
"""Convert The Swarm Within integrated TXT to PDF and DOCX"""

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY
import re

def create_docx(input_file, output_file):
    """Create DOCX from TXT file"""
    doc = Document()

    # Set margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    with open(input_file, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    for line in lines:
        line = line.rstrip()

        if not line:
            # Empty line
            doc.add_paragraph()
        elif line.startswith('The Swarm Within:'):
            # Title
            p = doc.add_paragraph(line)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.runs[0].bold = True
            p.runs[0].font.size = Pt(18)
        elif re.match(r'^[IVX]+\.\s+[A-Z]', line):
            # Main section header (I., II., III., etc.)
            p = doc.add_paragraph(line)
            p.runs[0].bold = True
            p.runs[0].font.size = Pt(14)
        elif re.match(r'^[A-Z]\.\s+[A-Z]', line):
            # Subsection header (A., B., C., etc.)
            p = doc.add_paragraph(line)
            p.runs[0].bold = True
            p.runs[0].font.size = Pt(12)
        elif line.isupper() and len(line) < 100:
            # ALL CAPS headers
            p = doc.add_paragraph(line)
            p.runs[0].bold = True
            p.runs[0].font.size = Pt(12)
        else:
            # Regular paragraph
            p = doc.add_paragraph(line)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.runs[0].font.size = Pt(11)

    doc.save(output_file)
    print(f"✓ Created {output_file}")

def create_pdf(input_file, output_file):
    """Create PDF from TXT file"""
    doc = SimpleDocTemplate(
        output_file,
        pagesize=letter,
        rightMargin=72,
        leftMargin=72,
        topMargin=72,
        bottomMargin=72
    )

    # Define styles
    styles = getSampleStyleSheet()

    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=18,
        alignment=TA_CENTER,
        spaceAfter=30,
        spaceBefore=12,
        fontName='Helvetica-Bold'
    )

    heading_style = ParagraphStyle(
        'CustomHeading',
        parent=styles['Heading2'],
        fontSize=14,
        spaceAfter=12,
        spaceBefore=12,
        fontName='Helvetica-Bold'
    )

    subheading_style = ParagraphStyle(
        'CustomSubHeading',
        parent=styles['Heading3'],
        fontSize=12,
        spaceAfter=10,
        spaceBefore=10,
        fontName='Helvetica-Bold'
    )

    body_style = ParagraphStyle(
        'CustomBody',
        parent=styles['BodyText'],
        fontSize=11,
        alignment=TA_JUSTIFY,
        spaceAfter=6,
        fontName='Times-Roman'
    )

    story = []

    with open(input_file, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    for line in lines:
        line = line.rstrip()

        if not line:
            story.append(Spacer(1, 0.2*inch))
        elif line.startswith('The Swarm Within:'):
            story.append(Paragraph(line, title_style))
        elif re.match(r'^[IVX]+\.\s+[A-Z]', line):
            story.append(Spacer(1, 0.15*inch))
            story.append(Paragraph(line, heading_style))
        elif re.match(r'^[A-Z]\.\s+[A-Z]', line):
            story.append(Paragraph(line, subheading_style))
        elif line.isupper() and len(line) < 100:
            story.append(Paragraph(line, subheading_style))
        else:
            # Escape special characters for reportlab
            line = line.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
            story.append(Paragraph(line, body_style))

    doc.build(story)
    print(f"✓ Created {output_file}")

if __name__ == '__main__':
    input_file = 'The_Swarm_Within_SLM_Agentic_Attacks_INTEGRATED.txt'

    print("Converting document...")
    create_docx(input_file, 'The_Swarm_Within_SLM_Agentic_Attacks_INTEGRATED.docx')
    create_pdf(input_file, 'The_Swarm_Within_SLM_Agentic_Attacks_INTEGRATED.pdf')
    print("\n✅ Conversion complete!")
