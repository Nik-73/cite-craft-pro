#!/usr/bin/env python3
"""Convert markdown research paper to DOCX and PDF formats."""

import re
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import markdown2

def markdown_to_docx(md_file, docx_file):
    """Convert markdown to DOCX with formatting."""
    # Read markdown file
    with open(md_file, 'r', encoding='utf-8') as f:
        content = f.read()

    # Create document
    doc = Document()

    # Set document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Parse markdown content
    lines = content.split('\n')

    for line in lines:
        line = line.strip()

        if not line:
            continue

        # Handle headers
        if line.startswith('# ') and not line.startswith('## '):
            # H1 - Title
            text = line.replace('# ', '')
            p = doc.add_paragraph()
            run = p.add_run(text)
            run.font.size = Pt(18)
            run.font.bold = True
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        elif line.startswith('## '):
            # H2 - Main sections
            text = line.replace('## ', '')
            p = doc.add_paragraph()
            run = p.add_run(text)
            run.font.size = Pt(14)
            run.font.bold = True

        elif line.startswith('### '):
            # H3 - Subsections
            text = line.replace('### ', '')
            p = doc.add_paragraph()
            run = p.add_run(text)
            run.font.size = Pt(12)
            run.font.bold = True

        elif line.startswith('**') and line.endswith('**'):
            # Bold headings
            text = line.replace('**', '')
            p = doc.add_paragraph()
            run = p.add_run(text)
            run.font.bold = True

        elif line.startswith('*') and line.endswith('*') and not line.startswith('**'):
            # Italic
            text = line.replace('*', '')
            p = doc.add_paragraph(text)

        elif line.startswith('-'):
            # List items
            text = line[1:].strip()
            # Process inline formatting
            text = process_inline_formatting(text)
            doc.add_paragraph(text, style='List Bullet')

        elif re.match(r'^[⁰¹²³⁴⁵⁶⁷⁸⁹]+', line):
            # Footnotes
            p = doc.add_paragraph(line)
            p.runs[0].font.size = Pt(10)

        else:
            # Regular paragraph
            # Process inline formatting
            text = process_inline_formatting(line)
            if text:
                doc.add_paragraph(text)

    # Save document
    doc.save(docx_file)
    print(f"DOCX file created: {docx_file}")

def process_inline_formatting(text):
    """Process inline markdown formatting."""
    # Remove complex markdown that doesn't translate well
    text = re.sub(r'\*supra\*', 'supra', text)
    text = re.sub(r'\*Id\.\*', 'Id.', text)
    return text

def markdown_to_html_pdf(md_file, html_file):
    """Convert markdown to HTML (intermediate step for PDF)."""
    # Read markdown file
    with open(md_file, 'r', encoding='utf-8') as f:
        content = f.read()

    # Convert to HTML
    html_content = markdown2.markdown(content, extras=['footnotes', 'tables'])

    # Create full HTML document with styling
    full_html = f"""
    <!DOCTYPE html>
    <html>
    <head>
        <meta charset="UTF-8">
        <style>
            body {{
                font-family: 'Times New Roman', serif;
                font-size: 12pt;
                line-height: 1.6;
                margin: 1in;
                max-width: 8.5in;
            }}
            h1 {{
                font-size: 18pt;
                font-weight: bold;
                text-align: center;
                margin-bottom: 20pt;
            }}
            h2 {{
                font-size: 14pt;
                font-weight: bold;
                margin-top: 20pt;
                margin-bottom: 10pt;
            }}
            h3 {{
                font-size: 12pt;
                font-weight: bold;
                margin-top: 15pt;
                margin-bottom: 8pt;
            }}
            p {{
                text-align: justify;
                margin-bottom: 10pt;
            }}
            em {{
                font-style: italic;
            }}
            strong {{
                font-weight: bold;
            }}
            sup {{
                font-size: 10pt;
            }}
        </style>
    </head>
    <body>
        {html_content}
    </body>
    </html>
    """

    # Save HTML
    with open(html_file, 'w', encoding='utf-8') as f:
        f.write(full_html)

    print(f"HTML file created: {html_file}")

if __name__ == '__main__':
    md_file = 'Tax_Dispute_Resolution_Research_Paper.md'
    docx_file = 'Tax_Dispute_Resolution_Research_Paper.docx'
    html_file = 'Tax_Dispute_Resolution_Research_Paper.html'

    print("Converting to DOCX...")
    markdown_to_docx(md_file, docx_file)

    print("\nConverting to HTML (for PDF conversion)...")
    markdown_to_html_pdf(md_file, html_file)

    print("\nConversion complete!")
    print(f"DOCX file: {docx_file}")
    print(f"HTML file: {html_file}")
    print("\nNote: For PDF conversion, you can:")
    print("1. Open the HTML file in a browser and print to PDF")
    print("2. Open the DOCX file in Microsoft Word and save as PDF")
