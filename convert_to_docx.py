#!/usr/bin/env python3
"""
Convert markdown research paper to editable Word DOCX format
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
import re

def setup_styles(doc):
    """Set up custom styles for the document"""

    # Normal style
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)
    style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    style.paragraph_format.first_line_indent = Inches(0.5)

    # Title style
    title_style = doc.styles['Title']
    title_style.font.name = 'Times New Roman'
    title_style.font.size = Pt(18)
    title_style.font.bold = True
    title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_style.paragraph_format.space_after = Pt(24)

    # Heading 1 (Roman numerals)
    h1_style = doc.styles['Heading 1']
    h1_style.font.name = 'Times New Roman'
    h1_style.font.size = Pt(14)
    h1_style.font.bold = True
    h1_style.font.all_caps = True
    h1_style.paragraph_format.space_before = Pt(18)
    h1_style.paragraph_format.space_after = Pt(12)
    h1_style.paragraph_format.first_line_indent = Inches(0)

    # Heading 2 (A, B, C)
    h2_style = doc.styles['Heading 2']
    h2_style.font.name = 'Times New Roman'
    h2_style.font.size = Pt(12)
    h2_style.font.bold = True
    h2_style.paragraph_format.space_before = Pt(14)
    h2_style.paragraph_format.space_after = Pt(10)
    h2_style.paragraph_format.first_line_indent = Inches(0)

    # Heading 3
    h3_style = doc.styles['Heading 3']
    h3_style.font.name = 'Times New Roman'
    h3_style.font.size = Pt(12)
    h3_style.font.bold = True
    h3_style.font.italic = True
    h3_style.paragraph_format.space_before = Pt(12)
    h3_style.paragraph_format.space_after = Pt(8)
    h3_style.paragraph_format.first_line_indent = Inches(0)

def parse_markdown_line(line):
    """Parse a line of markdown and return its type and content"""
    line = line.rstrip()

    if not line:
        return 'blank', ''

    # Check for headers
    if line.startswith('# '):
        return 'title', line[2:]
    elif line.startswith('## '):
        return 'h1', line[3:]
    elif line.startswith('### '):
        return 'h2', line[4:]
    elif line.startswith('#### '):
        return 'h3', line[5:]

    # Check for list items
    if re.match(r'^[\*\-] ', line):
        return 'list', line[2:]
    if re.match(r'^\d+\. ', line):
        match = re.match(r'^\d+\. (.*)', line)
        return 'list', match.group(1)

    # Regular paragraph
    return 'paragraph', line

def process_inline_formatting(paragraph, text):
    """Process inline formatting like bold, italic, and citations"""
    # Split by bold markers
    parts = re.split(r'(\*\*.*?\*\*)', text)

    for part in parts:
        if not part:
            continue

        if part.startswith('**') and part.endswith('**'):
            # Bold text
            content = part[2:-2]
            run = paragraph.add_run(content)
            run.bold = True
        else:
            # Regular text (may contain other formatting)
            # Handle italics
            italic_parts = re.split(r'(\*.*?\*)', part)
            for ipart in italic_parts:
                if not ipart:
                    continue

                if ipart.startswith('*') and ipart.endswith('*') and not ipart.startswith('**'):
                    content = ipart[1:-1]
                    run = paragraph.add_run(content)
                    run.italic = True
                else:
                    # Handle citations [Source X]
                    citation_parts = re.split(r'(\[Source \d+\])', ipart)
                    for cpart in citation_parts:
                        if not cpart:
                            continue

                        if cpart.startswith('[Source '):
                            run = paragraph.add_run(cpart)
                            run.font.superscript = True
                        else:
                            paragraph.add_run(cpart)

def convert_markdown_to_docx(md_file, citations_file, output_file):
    """Convert markdown files to Word document"""

    print(f"Reading {md_file}...")
    with open(md_file, 'r', encoding='utf-8') as f:
        md_lines = f.readlines()

    print(f"Reading {citations_file}...")
    with open(citations_file, 'r', encoding='utf-8') as f:
        citations_lines = f.readlines()

    print("Creating Word document...")
    doc = Document()

    # Set up page margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Set up styles
    setup_styles(doc)

    # Process main content
    print("Processing main content...")
    in_list = False

    for line in md_lines:
        line_type, content = parse_markdown_line(line)

        if line_type == 'blank':
            in_list = False
            continue

        elif line_type == 'title':
            p = doc.add_paragraph(content, style='Title')

        elif line_type == 'h1':
            p = doc.add_paragraph(content, style='Heading 1')
            in_list = False

        elif line_type == 'h2':
            p = doc.add_paragraph(content, style='Heading 2')
            in_list = False

        elif line_type == 'h3':
            p = doc.add_paragraph(content, style='Heading 3')
            in_list = False

        elif line_type == 'list':
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.first_line_indent = Inches(0)
            p.paragraph_format.left_indent = Inches(0.5)
            process_inline_formatting(p, content)
            in_list = True

        elif line_type == 'paragraph':
            p = doc.add_paragraph(style='Normal')
            process_inline_formatting(p, content)
            in_list = False

    # Add page break before citations
    doc.add_page_break()

    # Process citations
    print("Processing citations...")
    citation_title = doc.add_paragraph('BLUEBOOK CITATIONS', style='Heading 1')
    citation_title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for line in citations_lines:
        line = line.strip()
        if not line or line.startswith('#'):
            continue

        # Check if it's a source heading
        if line.startswith('**[Source'):
            # Extract source and add as hanging indent
            p = doc.add_paragraph(style='Normal')
            p.paragraph_format.left_indent = Inches(0.5)
            p.paragraph_format.first_line_indent = Inches(-0.5)
            p.paragraph_format.space_after = Pt(8)
            process_inline_formatting(p, line)
        elif line.startswith('**'):
            # Section heading in citations
            p = doc.add_paragraph(style='Heading 2')
            process_inline_formatting(p, line)
        elif line.startswith('##'):
            # Skip markdown headers
            continue
        elif line.startswith('---'):
            # Horizontal rule
            doc.add_paragraph('_' * 50)
        elif line:
            # Regular citation text
            p = doc.add_paragraph(style='Normal')
            p.paragraph_format.first_line_indent = Inches(0)
            p.paragraph_format.space_after = Pt(6)
            process_inline_formatting(p, line)

    # Save document
    print(f"Saving to {output_file}...")
    doc.save(output_file)
    print(f"✓ Word document created successfully: {output_file}")

def main():
    convert_markdown_to_docx(
        'refined_slm_research_paper.md',
        'bluebook_citations.md',
        'The_Swarm_Within_SLM_Agentic_Attacks.docx'
    )

if __name__ == '__main__':
    main()
